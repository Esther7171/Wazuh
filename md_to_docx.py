import tkinter as tk
from tkinter import filedialog, messagebox
from markdown import markdown
from bs4 import BeautifulSoup
from docx import Document
import os

def md_to_docx(md_path, docx_path):
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        html = markdown(md_content)
        soup = BeautifulSoup(html, 'html.parser')
        document = Document()

        for tag in soup.find_all(['h1', 'h2', 'h3', 'p', 'li']):
            text = tag.get_text()
            if tag.name.startswith('h'):
                level = int(tag.name[1])
                document.add_heading(text, level=level)
            elif tag.name == 'li':
                document.add_paragraph(text, style='List Bullet')
            else:
                document.add_paragraph(text)

        document.save(docx_path)
        messagebox.showinfo("Success", f"Document saved to:\n{docx_path}")
    except Exception as e:
        messagebox.showerror("Error", str(e))

def browse_md():
    file_path = filedialog.askopenfilename(
        filetypes=[("Markdown files", "*.md")], title="Select README.md"
    )
    if file_path:
        entry_md.delete(0, tk.END)
        entry_md.insert(0, file_path)

def convert_file():
    md_path = entry_md.get()
    if not os.path.isfile(md_path):
        messagebox.showerror("Error", "Invalid Markdown file selected.")
        return

    docx_path = filedialog.asksaveasfilename(
        defaultextension=".docx", filetypes=[("Word Document", "*.docx")], title="Save as"
    )
    if docx_path:
        md_to_docx(md_path, docx_path)

# GUI setup
root = tk.Tk()
root.title("Markdown to DOCX Converter")

frame = tk.Frame(root, padx=10, pady=10)
frame.pack()

label_md = tk.Label(frame, text="Select README.md file:")
label_md.grid(row=0, column=0, sticky="w")

entry_md = tk.Entry(frame, width=50)
entry_md.grid(row=1, column=0)

btn_browse = tk.Button(frame, text="Browse", command=browse_md)
btn_browse.grid(row=1, column=1, padx=5)

btn_convert = tk.Button(frame, text="Convert to DOCX", command=convert_file, bg="#4CAF50", fg="white")
btn_convert.grid(row=2, column=0, columnspan=2, pady=10)

root.mainloop()

# Requirement 
# pip install markdown python-docx beautifulsoup4
