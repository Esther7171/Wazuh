from docx import Document
from docx.shared import Pt

# Create a new Word document
doc = Document()

# Title
doc.add_heading('Comparison: Sachet SIEM + XDR vs Vision One XDR', level=1)

# Introduction about SIEM and XDR
doc.add_heading('What is SIEM and XDR?', level=2)
doc.add_paragraph(
    "SIEM (Security Information and Event Management) collects and aggregates log data from various sources "
    "across the network, analyzes and correlates events to detect suspicious activities, and supports compliance "
    "through centralized logging and historical data retention.\n"
    "XDR (Extended Detection and Response) focuses on real-time threat detection and automated response by "
    "integrating multiple security products such as endpoints, network, email, and cloud security to provide "
    "a unified threat detection and response platform."
)

# Table 1: Capabilities comparison
doc.add_heading('Capabilities Comparison', level=2)

table1 = doc.add_table(rows=1, cols=3)
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Capability'
hdr_cells[1].text = 'Sachet SIEM + XDR'
hdr_cells[2].text = 'Vision One XDR'

capabilities = [
    ("Collect logs from multiple devices (firewalls, servers, routers)", 
     "Yes, collects logs from wide range of devices and applications", 
     "Limited mainly to endpoints, network, email, and cloud telemetry"),
    ("Type of logs it can collect", 
     "System logs, application logs, firewall, IDS/IPS, endpoint logs, network flows", 
     "Endpoint telemetry, network traffic, email events, cloud logs"),
    ("Analyze and correlate data from various sources", 
     "Automatic rule-based correlation using decoders, CDB lists, and correlation rules", 
     "Automatic correlation across endpoints, network, email, and cloud data"),
    ("Detect suspicious behavior?", 
     "Yes, based on log analysis and correlation rules", 
     "Yes, via real-time telemetry and AI-driven analytics"),
    ("Automatically respond to threats", 
     "Manual or script-based responses using Active Response; supports custom scripts for blocking IPs, isolating endpoints, terminating processes; can integrate with external SOAR", 
     "Automated SOAR-like response (isolate endpoint, kill process, block IP) with built-in playbooks"),
    ("Visualize complete attack chain (who, where, how)", 
     "Basic visualization with alerts and reports", 
     "Advanced timeline views, attack chain visualization, and incident diagrams"),
    ("Support compliance/auditing", 
     "Built-in support for PCI-DSS, ISO 27001, HIPAA, GDPR, with centralized logs", 
     "Built-in compliance reports/templates for PCI, NIST, GDPR, HIPAA"),
    ("Retain historical log data", 
     "Yes, supports long-term storage for auditing and forensic analysis", 
     "Limited, focused on recent telemetry data"),
    ("Work across all vendors", 
     "Yes, supports multi-vendor log collection and integration", 
     "Primarily within Trend Micro ecosystem, supports 3rd party via APIs"),
    ("Integration", 
     "ELK Stack, Suricata, VirusTotal, Docker, MITRE ATT&CK, 3rd party APIs", 
     "Tight Trend Micro ecosystem plus 3rd party integrations via APIs"),
    ("Threat Intelligence", 
     "Real-time threat monitoring using VirusTotal, NVD database, hunting with CVEs, Elastest, YARA", 
     "Built-in global TI from Trend Micro Smart Protection Network"),
    ("Threat Correlation", 
     "Automatic correlation with custom rules, decoders, CDB lists", 
     "Automatic correlation across endpoints, network, email, and cloud"),
    ("Response Capabilities", 
     "Manual/script-based via Active Response; customizable automated scripts; integrates with external SOAR and orchestration tools", 
     "Automated SOAR workflows with built-in playbooks"),
    ("CVE detection in systems and software", 
     "Yes, scans logs and endpoint data for CVEs using integrated threat intel", 
     "Yes, via integrated vulnerability and threat intelligence feeds"),
    ("Forensics / Investigation", 
     "Provides basic logs and alerts with manual analysis; supports generating detailed reports and diagrams", 
     "Advanced timeline views, attack chains, incident visualization with graphical diagrams"),
    ("SOAR Capabilities", 
     "Limited; supports scripts and manual integrations; can save workflows", 
     "Built-in or integrated SOAR with saved automated workflows"),
    ("Compliance Support", 
     "Built-in system for PCI-DSS, ISO 27001, HIPAA, GDPR", 
     "Built-in reports and templates for PCI, NIST, GDPR, HIPAA")
]

for cap, sachet, vision in capabilities:
    row_cells = table1.add_row().cells
    row_cells[0].text = cap
    row_cells[1].text = sachet
    row_cells[2].text = vision

# Table 2: FIM + YARA vs DLP comparison
doc.add_heading('File Integrity Monitoring (FIM) + YARA vs Data Loss Prevention (DLP)', level=2)

table2 = doc.add_table(rows=1, cols=3)
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Feature'
hdr_cells[1].text = 'Sachet FIM + YARA'
hdr_cells[2].text = 'Vision One DLP'

fim_dlp_features = [
    ("File change detection", "Monitors file create, modify, delete, rename, permission changes", "Monitors and controls file transfers and access"),
    ("File version tracking", "Tracks file versions and changes over time", "Limited version control capabilities"),
    ("Malware obfuscation detection", "Detects obfuscated malware using YARA signatures", "Focuses more on data leakage than malware detection"),
    ("Detect leveraging attacks", "Detects suspicious file actions indicating lateral movement or privilege abuse", "Detects unauthorized data exfiltration attempts"),
    ("Detect lateral movement", "Monitors for suspicious file and process behavior indicative of lateral movement", "Limited to data loss and leakage prevention"),
    ("Real-time alerts", "Generates real-time alerts on suspicious file changes and YARA hits", "Generates alerts on data leakage policy violations"),
]

for feat, sachet, vision in fim_dlp_features:
    row_cells = table2.add_row().cells
    row_cells[0].text = feat
    row_cells[1].text = sachet
    row_cells[2].text = vision

# Save the document
file_path = "/mnt/data/Sachet_vs_VisionOne_Comparison.docx"
doc.save(file_path)

file_path
